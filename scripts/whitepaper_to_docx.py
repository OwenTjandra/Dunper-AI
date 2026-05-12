#!/usr/bin/env python3
"""
Convert docs/whitepaper.md → docs/whitepaper.docx

Handles the subset of Markdown the whitepaper actually uses:
- Headings (#, ##, ###)
- Paragraphs with inline **bold**, *italic*, `code`, [text](url)
- Bullet lists (- ...)
- Tables ( | a | b | )
- Code blocks (``` ... ``` — rendered as monospace)
- Block quotes (> ...)
- Horizontal rules (---)
- Italic emphasis blocks
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC  = ROOT / "docs" / "whitepaper.md"
DST  = ROOT / "docs" / "whitepaper.docx"

BRAND_NAVY  = RGBColor(0x0A, 0x14, 0x30)
BRAND_BLUE  = RGBColor(0x1E, 0x3A, 0x8A)
BODY_GREY   = RGBColor(0x47, 0x55, 0x69)
MUTED_GREY  = RGBColor(0x94, 0xA3, 0xB8)

# ───────────────────────────────────────────────────────────── helpers ──

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"           # **bold**
    r"|(\*[^*]+\*)"              # *italic*
    r"|(`[^`]+`)"                # `code`
    r"|(\[[^\]]+\]\([^)]+\))"    # [text](url)
)


def add_inline(paragraph, text, *, base_font="Calibri", base_size=11, color=BRAND_NAVY):
    """Append text to a paragraph, honoring inline **bold**, *italic*, `code`, [link](url)."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.font.color.rgb = color
        tok = m.group(0)
        if tok.startswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
            run.font.color.rgb = BRAND_NAVY
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = BRAND_BLUE
        elif tok.startswith("["):
            mtxt = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if mtxt:
                run = paragraph.add_run(mtxt.group(1))
                run.font.color.rgb = BRAND_BLUE
                run.underline = True
            else:
                run = paragraph.add_run(tok)
        elif tok.startswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
            run.font.color.rgb = color
        run.font.name = run.font.name or base_font
        run.font.size = run.font.size or Pt(base_size)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = Pt(base_size)
        run.font.color.rgb = color


def set_cell_shading(cell, fill_hex):
    """Apply a fill color to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(24)
        run.font.color.rgb = BRAND_NAVY
    elif level == 2:
        run.font.size = Pt(18)
        run.font.color.rgb = BRAND_NAVY
    else:
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_BLUE


def add_paragraph(doc, text, *, italic=False, size=11, color=BRAND_NAVY, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    if italic:
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Calibri"
    else:
        add_inline(p, text, base_size=size, color=color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text, base_size=11)


def add_blockquote(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    run_first = True
    for line in lines:
        if not run_first:
            p.add_run("\n")
        add_inline(p, line, base_size=11, color=BODY_GREY)
        run_first = False
    # italicize the whole paragraph
    for r in p.runs:
        r.italic = True


def add_codeblock(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for i, line in enumerate(lines):
        run = p.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        run.font.name = "Menlo"
        run.font.size = Pt(9.5)
        run.font.color.rgb = BRAND_NAVY


def add_table(doc, header_cells, body_rows):
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    # Header
    hdr = table.rows[0]
    for j, h in enumerate(header_cells):
        cell = hdr.cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.paragraph_format.space_after = Pt(0)
        set_cell_shading(cell, "0A1430")
    # Body
    for i, row in enumerate(body_rows):
        cells = table.rows[i + 1].cells
        for j, v in enumerate(row):
            if j >= len(cells):
                continue
            c = cells[j]
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, v.strip(), base_size=10.5)
            if i % 2 == 1:
                set_cell_shading(c, "F8FAFC")


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "94A3B8")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ───────────────────────────────────────────────────────── parser/renderer ──


def build():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # Body default
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            j = i + 1
            block = []
            while j < n and not lines[j].strip().startswith("```"):
                block.append(lines[j])
                j += 1
            add_codeblock(doc, block)
            i = j + 1
            continue

        # Table — header line followed by --- separator
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < n and re.match(r"^\|\s*[-:|\s]+\|\s*$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            j = i + 2
            rows = []
            while j < n and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            add_table(doc, header, rows)
            doc.add_paragraph()  # spacer after table
            i = j
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            add_hr(doc)
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*-\s+", line):
            # gather contiguous bullets
            while i < n and re.match(r"^\s*-\s+", lines[i]):
                add_bullet(doc, re.sub(r"^\s*-\s+", "", lines[i]).rstrip())
                i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            block = []
            while i < n and lines[i].strip().startswith("> "):
                block.append(lines[i].strip()[2:])
                i += 1
            add_blockquote(doc, block)
            continue

        # Standalone italic line at the bottom (`*...*`)
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            add_paragraph(doc, stripped.strip("*"), italic=True, color=MUTED_GREY, size=9.5)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph — accumulate consecutive non-empty non-special lines
        para_lines = [line.rstrip()]
        j = i + 1
        while j < n:
            nxt = lines[j].rstrip()
            ns  = nxt.strip()
            if not ns:
                break
            if re.match(r"^(#{1,3})\s+", ns):
                break
            if ns.startswith("```") or ns.startswith("> ") or ns.startswith("- "):
                break
            if ns.startswith("|") and ns.endswith("|"):
                break
            if ns == "---":
                break
            para_lines.append(nxt)
            j += 1
        add_paragraph(doc, " ".join(para_lines), space_after=8)
        i = j

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    return DST


if __name__ == "__main__":
    out = build()
    print(f"wrote {out}")
